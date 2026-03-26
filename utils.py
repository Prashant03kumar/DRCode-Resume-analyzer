import io
import re
import pypdf
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from config import logger

# ─────────────────────────────────────────────────────────────────────────────
# Text Extraction
# ─────────────────────────────────────────────────────────────────────────────

async def extract_text_from_file(file, file_name: str) -> str:
    """Extracts text from PDF, DOCX, or TXT formats."""
    file_bytes = await file.download_as_bytearray()
    text = ""
    file_ext = file_name.split(".")[-1].lower()

    try:
        if file_ext == "pdf":
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        elif file_ext in ["docx", "doc"]:
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif file_ext == "txt":
            text = file_bytes.decode("utf-8")
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        logger.error(f"Error parsing {file_name}: {e}")
        return None

    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# Resume Parsing Helper
# ─────────────────────────────────────────────────────────────────────────────

def parse_resume_sections(resume_text: str):
    """
    Parses resume plain text into structured sections.
    Returns list of dicts: {'type': 'name'|'contact'|'section_header'|'bullet'|'text', 'content': str}
    """
    lines = resume_text.split('\n')
    sections = []
    first_non_empty = True

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            sections.append({'type': 'spacer', 'content': ''})
            continue

        # First non-empty line → name
        if first_non_empty:
            sections.append({'type': 'name', 'content': line})
            first_non_empty = False
            continue

        # Contact info line (contains @ or phone patterns or |)
        if re.search(r'@|\||\+?\d[\d\s\-]{7,}|linkedin|github|http', line, re.IGNORECASE):
            sections.append({'type': 'contact', 'content': line})
            continue

        # Section header: ALL CAPS line (length 3-40) OR ends with : (short)
        stripped_colon = line.rstrip(':').strip()
        if (
            (line.isupper() and 3 <= len(line) <= 50)
            or (line.endswith(':') and len(stripped_colon) <= 35 and stripped_colon.isupper())
        ):
            sections.append({'type': 'section_header', 'content': stripped_colon})
            continue

        # Bullet point
        if line.startswith(('•', '-', '*', '–', '▪', '◦')):
            sections.append({'type': 'bullet', 'content': line.lstrip('•-*–▪◦ ').strip()})
            continue

        # Default: normal text
        sections.append({'type': 'text', 'content': line})

    return sections


# ─────────────────────────────────────────────────────────────────────────────
# PDF Generator (looks like an actual CV)
# ─────────────────────────────────────────────────────────────────────────────

ACCENT_COLOR = colors.HexColor('#1b4f72')   # Deep blue
LIGHT_GRAY   = colors.HexColor('#f2f3f4')
LINE_COLOR   = colors.HexColor('#2e86c1')
TEXT_COLOR   = colors.HexColor('#1a1a1a')
SUB_COLOR    = colors.HexColor('#555555')


def generate_pdf(resume_text: str, original_filename: str) -> tuple[io.BytesIO, str]:
    """
    Generates a professional, CV-styled PDF from the rewritten resume text.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.8 * cm, leftMargin=1.8 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )

    # ── Define Styles ──────────────────────────────────────────────────────
    name_style = ParagraphStyle(
        'CandidateName',
        fontSize=22, fontName='Helvetica-Bold',
        textColor=ACCENT_COLOR, alignment=TA_CENTER,
        spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        'ContactInfo',
        fontSize=9, fontName='Helvetica',
        textColor=SUB_COLOR, alignment=TA_CENTER,
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        'SectionHeader',
        fontSize=11, fontName='Helvetica-Bold',
        textColor=ACCENT_COLOR,
        spaceBefore=10, spaceAfter=3,
        leftIndent=0,
    )
    text_style = ParagraphStyle(
        'BodyText',
        fontSize=9.5, fontName='Helvetica',
        textColor=TEXT_COLOR,
        leading=14, spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        'BulletItem',
        fontSize=9.5, fontName='Helvetica',
        textColor=TEXT_COLOR,
        leading=14, spaceAfter=2,
        leftIndent=14, bulletIndent=4,
    )

    # ── Build Story ────────────────────────────────────────────────────────
    story = []
    sections = parse_resume_sections(resume_text)

    for item in sections:
        kind = item['type']
        text = item['content']

        def safe(t):
            return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        if kind == 'name':
            story.append(Paragraph(safe(text), name_style))

        elif kind == 'contact':
            story.append(Paragraph(safe(text), contact_style))

        elif kind == 'section_header':
            story.append(Spacer(1, 4))
            story.append(HRFlowable(
                width='100%', thickness=1.5,
                color=ACCENT_COLOR, spaceAfter=4,
            ))
            story.append(Paragraph(safe(text), section_style))

        elif kind == 'bullet':
            story.append(Paragraph(f"• &nbsp;{safe(text)}", bullet_style))

        elif kind == 'text':
            story.append(Paragraph(safe(text), text_style))

        elif kind == 'spacer':
            story.append(Spacer(1, 4))

    doc.build(story)
    buf.seek(0)
    base = original_filename.rsplit('.', 1)[0]
    return buf, f"{base}_DrCode_Optimized.pdf"


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Generator (looks like an actual CV)
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(resume_text: str, original_filename: str) -> tuple[io.BytesIO, str]:
    """Generates a professional DOCX CV from the rewritten resume text."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    sections = parse_resume_sections(resume_text)
    DARK_BLUE = RGBColor(0x1b, 0x4f, 0x72)
    GRAY      = RGBColor(0x55, 0x55, 0x55)

    for item in sections:
        kind = item['type']
        text = item['content']

        if kind == 'name':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_after = Pt(2)

        elif kind == 'contact':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            p.paragraph_format.space_after = Pt(6)

        elif kind == 'section_header':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE
            # Add bottom border to simulate a horizontal rule
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1b4f72')
            pBdr.append(bottom)
            pPr.append(pBdr)

        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)

        elif kind == 'text':
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)

        elif kind == 'spacer':
            p = doc.add_paragraph('')
            p.paragraph_format.space_after = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    base = original_filename.rsplit('.', 1)[0]
    return buf, f"{base}_DrCode_Optimized.docx"


# ─────────────────────────────────────────────────────────────────────────────
# TXT Generator
# ─────────────────────────────────────────────────────────────────────────────

def generate_txt(resume_text: str, original_filename: str) -> tuple[io.BytesIO, str]:
    """Returns the plain text resume as a downloadable TXT file."""
    buf = io.BytesIO(resume_text.encode('utf-8'))
    buf.seek(0)
    base = original_filename.rsplit('.', 1)[0]
    return buf, f"{base}_DrCode_Optimized.txt"
